from docx import Document

doc = Document("Efficiency and Accuracy of Few.docx")

# Replace incorrect data in Paragraph 27
p27 = doc.paragraphs[27]
p27.text = "The basic CNN model achieved lower baseline performance. Adding Attention significantly improved the accuracy to 99.64%, meaning the model became better at focusing on important regions and features in the data."

# Replace incorrect data in Paragraph 28
p28 = doc.paragraphs[28]
p28.text = "CNN+Mamba achieved an optimal 99.93%, showing that Mamba improved sequence and contextual understanding to near perfection. The combined CNN+Attention+Mamba model also achieved high performance, but CNN+Mamba alone proved to be the most optimal and stable architecture."

doc.save("Efficiency and Accuracy of Few_Fixed.docx")
print("Saved as Efficiency and Accuracy of Few_Fixed.docx")
