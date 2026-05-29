import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_bar_chart():
    # Data
    models = [
        'ProtoNet\n(CNN+Mamba)', 
        'RelationNet\n(CNN+Mamba)', 
        'ProtoNet\n(CNN+Attn)', 
        'RelationNet\n(CNN+Attn)', 
        'MAML\n(CNN)'
    ]
    accuracies = [99.93, 99.79, 99.64, 99.29, 50.00]
    
    fig, ax = plt.subplots(figsize=(8, 5))
    bars = ax.bar(models, accuracies, color=['#4C72B0', '#55A868', '#C44E52', '#8172B2', '#CCB974'])
    
    ax.set_ylabel('Accuracy (%)')
    ax.set_title('Few-Shot Learning Framework Accuracies')
    ax.set_ylim(0, 110)
    
    # Add data labels
    for bar in bars:
        yval = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, yval + 1, f'{yval}%', ha='center', va='bottom', fontweight='bold')
    
    plt.tight_layout()
    chart_path = 'accuracy_chart.png'
    plt.savefig(chart_path, dpi=300)
    plt.close()
    return chart_path

def set_two_column(doc):
    # Set the whole document to 2 columns
    section = doc.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')  # 0.5 inch spacing between columns

def main():
    chart_path = create_bar_chart()
    doc = Document()
    
    # Title - usually spans both columns, but we will add it normally
    title = doc.add_heading('Efficiency and Accuracy of Few-Shot Learning Frameworks', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Apply 2-column layout to the section
    set_two_column(doc)
    
    # 1. Abstract / Overview
    doc.add_heading('1. Abstract', level=1)
    doc.add_paragraph(
        "This document presents a pure performance comparison of various Few-Shot Learning (FSL) "
        "frameworks applied to hyperspectral crop signature classification. The objective "
        "is to evaluate the efficiency and accuracy of different combinations of FSL frameworks and "
        "backbone architectures."
    )
    
    # 2. Evaluated Frameworks
    doc.add_heading('2. Frameworks Evaluated', level=1)
    doc.add_paragraph("The evaluation was conducted across three state-of-the-art FSL frameworks:")
    doc.add_paragraph("Prototypical Networks (ProtoNet): Metric-based FSL representing classes by average embeddings.", style='List Bullet')
    doc.add_paragraph("Relation Networks (RelationNet): Metric-based FSL learning a non-linear similarity metric.", style='List Bullet')
    doc.add_paragraph("Model-Agnostic Meta-Learning (MAML): Optimization-based FSL for fast adaptation.", style='List Bullet')
    
    # 3. Performance Graph
    doc.add_heading('3. Accuracy Comparison', level=1)
    doc.add_paragraph("The following graph and data summarize the final validation accuracies achieved after optimization.")
    
    # Insert Graph (Width constrained to fit inside a column)
    doc.add_picture('fsl_protonet_best_plot.png', width=Inches(3.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 1: ProtoNet Accuracy", style='Caption')
    
    doc.add_picture('fsl_relationnet_best_plot.png', width=Inches(3.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 2: RelationNet Accuracy", style='Caption')
    
    # 4. Results Data
    doc.add_heading('4. Results Data', level=1)
    
    records = [
        ('ProtoNet + CNN+Mamba', '99.93%'),
        ('RelationNet + CNN+Mamba', '99.79%'),
        ('ProtoNet + CNN+Attn', '99.64%'),
        ('RelationNet + CNN+Attn', '99.29%'),
        ('MAML + CNN', '~50.00%')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Architecture'
    hdr_cells[1].text = 'Accuracy'
    
    for arch, acc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = arch
        row_cells[1].text = acc
        
    # 5. Key Findings
    doc.add_heading('5. Key Findings', level=1)
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run('Optimal Architecture: ').bold = True
    p1.add_run("ProtoNet with CNN+Mamba achieved near-perfect accuracy (99.93%). Mamba's efficient state-space modeling handles sequential hyperspectral data exceptionally well.")
    
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run('MAML Limitations: ').bold = True
    p2.add_run("MAML suffered from second-order gradient explosion when paired with deep backbones like CNN+Mamba on this timeseries dataset, limiting accuracy to ~50%.")
    
    # Save the updated docx
    output_filename = 'FSL_Accuracy_Report_TwoColumn.docx'
    doc.save(output_filename)
    print(f"Successfully generated {output_filename} with 2-column layout and graphs.")

if __name__ == '__main__':
    main()
