from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_two_column(doc):
    section = doc.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')

doc = Document()

# Title
title = doc.add_heading('Efficiency and Accuracy of Few-Shot Learning Frameworks', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

set_two_column(doc)

# Abstract
doc.add_heading('Abstract', level=1)
doc.add_paragraph("This document presents a pure performance comparison of various Few-Shot Learning (FSL) frameworks applied to hyperspectral crop signature classification. The objective is to evaluate the efficiency and accuracy of different combinations of FSL frameworks and backbone architectures.")

# Frameworks Evaluated
doc.add_heading('Frameworks Evaluated', level=1)
doc.add_paragraph("The evaluation was conducted across three FSL frameworks:")
doc.add_paragraph("Prototypical Networks (ProtoNet): Metric-based FSL representing classes by average embeddings.", style='List Bullet')
doc.add_paragraph("Relation Networks (RelationNet): Metric-based FSL learning a non-linear similarity metric.", style='List Bullet')
doc.add_paragraph("Model-Agnostic Meta-Learning (MAML): Optimization-based FSL for fast adaptation.", style='List Bullet')

doc.add_heading('Accuracy Comparison', level=1)

# 1 ) RelationNet
doc.add_heading('1) RelationNet', level=2)
doc.add_paragraph("The RelationNet architectures achieved consistently high accuracies across trials.")
doc.add_paragraph("CNN+Attention and CNN+Mamba both improved the model performance compared to the basic CNN setup. CNN+Mamba especially showed strong stability and high median accuracy.")
doc.add_paragraph("The combined CNN+Attention+Mamba version achieved the best overall performance, showing that integrating multiple feature-learning mechanisms improved the relational learning capability of RelationNet.")

# Add RelationNet Graph
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    p.add_run().add_picture('fsl_relationnet_best_plot.png', width=Inches(3.0))
    doc.add_paragraph("Figure 1: RelationNet Curves", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_heading('Loss & Accuracy Curves (RelationNet)', level=3)
doc.add_paragraph("All RelationNet models reduced loss steadily during training. However, CNN+Mamba and CNN+Attention+Mamba reduced loss faster and more smoothly than the others.")
doc.add_paragraph("This suggests that Mamba likely helped capture broader contextual dependencies while Attention enhanced important feature selection.")
doc.add_paragraph("The accuracy curves show that all models improved over time, but the hybrid models improved faster and achieved higher final accuracies.")
doc.add_paragraph("Even though the CNN+Mamba and CNN+Attention+Mamba models began with slightly lower accuracy initially, they quickly caught up and eventually outperformed the basic CNN model. Toward the end of training, their accuracies stabilized near 100%, indicating excellent learning and generalization performance.")

# 2 ) ProtoNet
doc.add_heading('2) ProtoNet', level=2)
doc.add_paragraph("The basic CNN model achieved lower baseline performance. Adding Attention significantly improved the accuracy to 99.64%, meaning the model became better at focusing on important regions and features in the data.")
doc.add_paragraph("CNN+Mamba achieved an optimal 99.93%, showing that Mamba improved sequence and contextual understanding to near perfection. The combined CNN+Attention+Mamba model also achieved high performance, but CNN+Mamba alone proved to be the most optimal and stable architecture.")

# Add ProtoNet Holdout Graph
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    p.add_run().add_picture('fsl_protonet_best_holdout.png', width=Inches(3.0))
    doc.add_paragraph("Figure 2: ProtoNet 5-Shot Holdout Distribution", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_paragraph("Even though some models started with lower accuracies initially, they improved steadily during training and became more stable toward the end. The combined architecture benefited from both attention-based feature selection and Mamba-based sequence learning, leading to stronger generalization.")

# Add ProtoNet Curves Graph
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    p.add_run().add_picture('fsl_protonet_best_curves.png', width=Inches(3.0))
    doc.add_paragraph("Figure 3: ProtoNet Loss & Accuracy Curves", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_heading('Loss & Accuracy (ProtoNet)', level=3)
doc.add_paragraph("The CNN-only model reduced loss slowly and had more fluctuations, indicating slower learning and less stable optimization. CNN+Attention initially reduced the loss quickly, but later showed some instability before improving again.")
doc.add_paragraph("CNN+Mamba reduced the loss much faster after the early episodes, showing efficient learning capability. The CNN+Attention+Mamba model had one of the smoothest and fastest reductions in loss, eventually reaching the lowest loss values.")
doc.add_paragraph("This indicates that combining Attention and Mamba helped the model converge faster and learn more meaningful representations.")
doc.add_paragraph("The CNN model improved gradually but saturated below the hybrid models. CNN+Attention started strong and improved steadily. CNN+Mamba had a slower beginning, but after several episodes its accuracy increased rapidly and eventually approached nearly 100% accuracy.")

# 3 ) MAML
doc.add_heading('3) MAML', level=2)
doc.add_paragraph("The MAML-based model showed lower overall performance compared to ProtoNet and RelationNet.")

# Add MAML Graph
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    p.add_run().add_picture('fsl_maml_best_plot.png', width=Inches(3.0))
    doc.add_paragraph("Figure 4: MAML Loss & Accuracy", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_paragraph("Although the CNN backbone allowed some learning, the model struggled to maintain stable generalization across trials. The lower median accuracy and wider spread indicate inconsistent adaptation performance in the few-shot setting.")
doc.add_paragraph("This suggests that MAML may require stronger feature extractors or additional optimization strategies to compete with the hybrid ProtoNet and RelationNet architectures.")

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph("The ProtoNet with CNN+Mamba achieved the highest accuracy of 99.93%, showing that the Mamba block helped the model learn feature relationships very effectively. RelationNet with CNN+Mamba also performed extremely well with 99.79% accuracy, proving that combining CNN with Mamba improves classification performance.")
doc.add_paragraph("The CNN+Attention versions also gave strong results, but slightly lower than the Mamba-based models. This suggests that attention mechanisms helped the model focus on important features, but Mamba provided better long-range feature understanding and stability.")
doc.add_paragraph("MAML with only CNN achieved around 50% accuracy, which is much lower than the other methods. This indicates that the basic CNN backbone struggled to generalize effectively in the few-shot setting compared to the hybrid architectures.")

# Add Accuracy Bar Chart
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    p.add_run().add_picture('accuracy_chart.png', width=Inches(3.0))
    doc.add_paragraph("Figure 5: Final Accuracy Comparison", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

from docx.shared import RGBColor

# Apply black color to all text and headings
for p in doc.paragraphs:
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

doc.save("Efficiency_and_Accuracy_Final_Black.docx")
print("Saved as Efficiency_and_Accuracy_Final_Black.docx")
