from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Few-Shot Learning (FSL) Architecture Comparison Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        "This document presents a pure performance comparison of various Few-Shot Learning (FSL) "
        "frameworks applied to our hyperspectral crop signature classification project. The objective "
        "is to evaluate the efficiency and accuracy of different combinations of FSL frameworks and "
        "backbone architectures."
    )
    
    # Frameworks Evaluated
    doc.add_heading('2. Frameworks Evaluated', level=1)
    doc.add_paragraph("The evaluation was conducted across three state-of-the-art FSL frameworks:")
    doc.add_paragraph("Prototypical Networks (ProtoNet): Metric-based FSL representing classes by average embeddings.", style='List Bullet')
    doc.add_paragraph("Relation Networks (RelationNet): Metric-based FSL learning a non-linear similarity metric via a neural network relation module.", style='List Bullet')
    doc.add_paragraph("Model-Agnostic Meta-Learning (MAML): Optimization-based FSL for fast adaptation.", style='List Bullet')
    
    # Performance Results
    doc.add_heading('3. Performance Results', level=1)
    doc.add_paragraph("The following table summarizes the final validation accuracies achieved by each framework and backbone architecture combination after rigorous optimization.")
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'FSL Framework'
    hdr_cells[1].text = 'Backbone Architecture'
    hdr_cells[2].text = 'Final Accuracy'
    hdr_cells[3].text = 'Status / Stability'
    
    records = [
        ('ProtoNet', 'CNN + Mamba', '99.93%', 'Optimal / Highly Stable'),
        ('RelationNet', 'CNN + Mamba', '99.79%', 'Verified / Stable'),
        ('ProtoNet', 'CNN + Attention', '99.64%', 'Verified / Stable'),
        ('RelationNet', 'CNN + Attention', '99.29%', 'Verified / Stable'),
        ('MAML', 'CNN (Original)', '~50.00%', 'Unstable (Gradient Explosion)')
    ]
    
    for fw, arch, acc, stat in records:
        row_cells = table.add_row().cells
        row_cells[0].text = fw
        row_cells[1].text = arch
        row_cells[2].text = acc
        row_cells[3].text = stat
        
    # Key Findings
    doc.add_heading('4. Key Findings', level=1)
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run('Optimal Architecture: ').bold = True
    p1.add_run("The combination of ProtoNet with a CNN + Mamba backbone achieved near-perfect accuracy (99.93%). Mamba's efficient state-space modeling handles the sequential hyperspectral data exceptionally well without overfitting.")
    
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run('Metric-Learning vs. Meta-Learning: ').bold = True
    p2.add_run("Metric-learning approaches (ProtoNet, RelationNet) drastically outperformed the optimization-based meta-learning approach (MAML).")
    
    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run('MAML Limitations: ').bold = True
    p3.add_run("MAML suffered from mathematical incompatibility (second-order gradient explosion) when paired with deep, complex backbones like CNN+Mamba on this specific timeseries dataset, limiting its accuracy to ~50%.")
    
    # Efficiency Analysis
    doc.add_heading('5. Efficiency Analysis', level=1)
    doc.add_paragraph("ProtoNet provides the highest computational efficiency during inference, as it only requires computing Euclidean distances to static class prototypes.", style='List Bullet')
    doc.add_paragraph("RelationNet requires slightly more computational overhead during inference due to the forward pass through the relation module for every query-support pair.", style='List Bullet')
    doc.add_paragraph("CNN + Mamba backbones processed the 1D hyperspectral signals efficiently without the memory bottlenecks typically associated with standard self-attention mechanisms.", style='List Bullet')
    
    doc.save('FSL_Accuracy_Report.docx')

if __name__ == '__main__':
    main()
